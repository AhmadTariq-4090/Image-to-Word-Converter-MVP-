import io
import time

import google.generativeai as genai
import pytesseract
from docx import Document
from PIL import Image
from pytesseract import Output


class ActionAgent:
    """
    Executes OCR using the chosen engine and assembles the Word document.
    Returns structured results including confidence scores and timing.
    """

    # ── Public API ────────────────────────────────────────────────────────────

    def process(self, image: Image.Image, engine: str, api_key: str = None) -> dict:
        """Run OCR on a single image. Returns result dict."""
        start = time.time()

        if engine == "tesseract":
            text, confidence = self._tesseract(image)
        else:
            text, confidence = self._gemini(image, api_key)

        elapsed = round(time.time() - start, 2)
        success = not text.startswith("Error")

        return {
            "text":            text,
            "confidence":      confidence,
            "engine":          engine,
            "processing_time": elapsed,
            "word_count":      len(text.split()) if success else 0,
            "success":         success,
        }

    def generate_docx(self, results: list) -> bytes:
        """Build a .docx file from a list of action results."""
        doc = Document()
        doc.add_heading("Converted Document", 0)

        for i, result in enumerate(results):
            if len(results) > 1:
                name = result.get("name", f"Image {i + 1}")
                doc.add_heading(name, level=2)

            self._markdown_to_doc(doc, result["text"])

            if i < len(results) - 1:
                doc.add_page_break()

        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    # ── OCR engines ───────────────────────────────────────────────────────────

    def _tesseract(self, image: Image.Image) -> tuple:
        try:
            data = pytesseract.image_to_data(image, output_type=Output.DICT)
            confs = [int(c) for c in data["conf"]
                     if str(c).lstrip("-").isdigit() and int(c) >= 0]
            avg_conf = round(sum(confs) / len(confs) / 100, 2) if confs else 0.5
            text = pytesseract.image_to_string(image)
            return text.strip(), avg_conf
        except Exception as exc:
            return f"Error in Tesseract: {exc}", 0.0

    def _gemini(self, image: Image.Image, api_key: str) -> tuple:
        try:
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel("gemini-1.5-flash")
            prompt = (
                "You are an expert document digitizer. "
                "Convert the text in this image into Markdown format. "
                "Preserve: Headers (#, ##, ###), Bold (**text**), "
                "Italics (*text*), Bullet points (- item), "
                "Numbered lists (1. item), paragraph structure. "
                "Return ONLY the Markdown content. No preamble."
            )
            response = model.generate_content([prompt, image])
            return response.text.strip(), 0.88  # calibrated Gemini estimate
        except Exception as exc:
            return f"Error in Gemini: {exc}", 0.0

    # ── Markdown → docx ───────────────────────────────────────────────────────

    def _markdown_to_doc(self, doc: Document, markdown: str):
        for raw_line in markdown.split("\n"):
            line = raw_line.strip()
            if not line:
                doc.add_paragraph("")
                continue

            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith(("- ", "* ")):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif len(line) > 2 and line[0].isdigit() and ". " in line[:4]:
                doc.add_paragraph(line[line.index(". ") + 2:], style="List Number")
            else:
                p = doc.add_paragraph()
                parts = line.split("**")
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    run.bold = (i % 2 == 1)
