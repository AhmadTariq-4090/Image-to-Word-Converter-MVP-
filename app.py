import streamlit as st
from PIL import Image
import pytesseract
from docx import Document
import io
import os
import google.generativeai as genai

# Page Config
st.set_page_config(page_title="Image-to-Word Converter", layout="wide")

# Title
st.title("📄 Image-to-Word Converter MVP")
st.markdown("Convert your images to editable Word documents using **OCR** or **AI Vision**.")

# --- Sidebar ---
st.sidebar.header("⚙️ Settings")
engine = st.sidebar.radio(
    "Select Processing Engine",
    ("Tesseract OCR (Fast/Free)", "AI Vision (Google Gemini)")
)

api_key = None
if "AI Vision" in engine:
    # Try to get from environment variable first
    env_key = os.getenv("AIzaSyA4quSrLXH1bpMTFOtKysZZh0kzebkaSUc")
    api_key = st.sidebar.text_input("Enter Google Gemini API Key", value=env_key if env_key else "", type="password")
    if not api_key:
        st.sidebar.warning("⚠️ API Key is required for AI Vision.")
    else:
        st.sidebar.success("API Key provided!")

# --- Main Interface ---
tab1, tab2 = st.tabs(["📤 Upload Image", "📸 Camera Capture"])

image_files = []

with tab1:
    uploaded_files = st.file_uploader("Upload Images (JPG/PNG)", type=['png', 'jpg', 'jpeg'], accept_multiple_files=True)
    if uploaded_files:
        image_files.extend(uploaded_files)
        st.success(f"{len(uploaded_files)} images uploaded.")

with tab2:
    camera_image = st.camera_input("Take a picture")
    if camera_image:
        image_files.append(camera_image)
        st.success("Image captured from camera.")

# --- Helper Functions ---

def process_with_tesseract(image):
    """Extracts text using Tesseract OCR."""
    try:
        # Optional: Add pre-processing here (grayscale, threshold) if needed
        # gray = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY) # Requires opencv
        text = pytesseract.image_to_string(image)
        return text
    except Exception as e:
        return f"Error in Tesseract OCR: {str(e)}"

def process_with_gemini(image, api_key):
    """Extracts text and formatting using Google Gemini."""
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        prompt = """
        You are an expert document digitizer. 
        Convert the text in this image into Markdown format.
        Preserve the following formatting exactly as seen in the image:
        - Headers (#, ##, ###)
        - Bold (**text**)
        - Italics (*text*)
        - Bullet points (- item) and Numbered lists (1. item)
        - Paragraph structure.
        Do not add any preamble or sentences like "Here is the markdown". Just return the markdown content.
        """
        
        response = model.generate_content([prompt, image])
        return response.text
    except Exception as e:
        return f"Error in Gemini Vision: {str(e)}"

def add_markdown_to_doc(doc, markdown_text):
    """Parses simple Markdown and adds it to the python-docx Document."""
    lines = markdown_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph("") # Add empty line
            continue
            
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        # List Items (Basic support)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. '): # Simple Check for "1. "
            p = doc.add_paragraph(line[3:], style='List Number')
        else:
            # Normal Paragraph with Bold/Italic support
            p = doc.add_paragraph()
            # A simple parser for bold (**text**)
            parts = line.split('**')
            for i, part in enumerate(parts):
                runner = p.add_run(part)
                if i % 2 == 1: # Odd indices are between ** ** (bold)
                    runner.bold = True
            # Note regarding Italics: This creates a limitation where we can't nest bold/italics easily 
            # with this split method. This is a basic MVP parser. Prod usage would need a regex parser.

# --- Processing Section ---
if image_files:
    st.divider()
    st.subheader(f"Processing {len(image_files)} Image(s)")
    
    # Preview
    cols = st.columns(min(len(image_files), 3))
    for idx, img_file in enumerate(image_files[:3]): # Show max 3 previews
        with cols[idx]:
            st.image(img_file, caption=f"Image {idx+1}", use_container_width=True)
    if len(image_files) > 3:
        st.info(f"...and {len(image_files)-3} more.")

    if st.button("🚀 Convert to Word Document", type="primary"):
        if "AI Vision" in engine and not api_key:
            st.error("Please provide a Google Gemini API Key in the sidebar.")
        else:
            with st.spinner("Processing... Please wait."):
                # 1. Initialize Docx
                doc = Document()
                doc.add_heading('Converted Document', 0)

                # 2. Loop through images
                for i, img_file in enumerate(image_files):
                    img = Image.open(img_file)
                    st.write(f"Processing Image {i+1}/{len(image_files)}...")
                    
                    # 3. Apply Engine Logic
                    extracted_text = ""
                    if "Tesseract" in engine:
                        extracted_text = process_with_tesseract(img)
                        # Add raw text to doc
                        doc.add_paragraph(extracted_text)
                        doc.add_page_break()
                        
                    elif "AI Vision" in engine:
                        extracted_text = process_with_gemini(img, api_key)
                        # Parse Markdown to Doc
                        add_markdown_to_doc(doc, extracted_text)
                        doc.add_page_break()
                
                # 4. Save and Provide Download
                bio = io.BytesIO()
                doc.save(bio)
                
                st.success("✅ Conversion Complete!")
                st.download_button(
                    label="⬇️ Download Word Document",
                    data=bio.getvalue(),
                    file_name="converted_document.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
